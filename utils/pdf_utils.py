# utils/pdf_utils.py

import os
from docx import Document
from docx2pdf import convert
from flask import current_app
import logging
from typing import Optional, Union
from datetime import date, datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt # Added Pt for font size
from num2words import num2words # For converting numbers to words
import docx.oxml # For page number field
from docx.oxml.ns import qn # Import qn for qualified names
import pythoncom # Import for CoInitialize

# Get the logger for the current module
logger = logging.getLogger(__name__)

class TCGenerator:
    """Generates Transfer Certificates (TC) as Word (.docx) and PDF documents."""

    def __init__(self, template_path: Optional[str] = None):
        """
        Initializes the TCGenerator.
        Args:
            template_path (str, optional): Path to the Word TC template.
                                           Defaults to TC_TEMPLATE_PATH from app config.
        """
        self.template_path = template_path or current_app.config.get('TC_TEMPLATE_PATH')
        self.output_path_base = current_app.config.get('TC_OUTPUT_PATH')

        if not self.output_path_base:
            logger.error("TC_OUTPUT_PATH is not configured in the application.")
            raise ValueError("TC output path is not configured.")

        # Ensure output directory exists
        try:
            os.makedirs(self.output_path_base, exist_ok=True)
        except OSError as e:
            logger.error(f"Could not create TC output directory {self.output_path_base}: {e}")
            raise

    def _get_safe_filename_part(self, value: str) -> str:
        """Cleans a string for use in filenames by replacing invalid characters."""
        return str(value).replace('/', '_').replace('\\', '_')

    def generate_tc_files(self, student_data: dict, tc_data: dict) -> tuple[str, str]:
        """
        Generates TC files (.docx and .pdf) from data.

        Args:
            student_data (dict): Dictionary containing student information.
            tc_data (dict): Dictionary containing TC specific data.

        Returns:
            tuple[str, str]: A tuple containing the paths to the generated (docx_path, pdf_path).

        Raises:
            FileNotFoundError: If the template path is specified but the file doesn't exist.
            RuntimeError: For other errors during document generation or conversion.
        """
        docx_path = ""
        try:
            # Step 1: Create the Word Document
            if self.template_path and os.path.exists(self.template_path):
                doc = Document(self.template_path)
                section = doc.sections[0]
                section.page_width = Inches(8.5)
                section.page_height = Inches(14)
                # Optionally set margins as above
            elif self.template_path:
                logger.error(f"TC template specified but not found: {self.template_path}")
                raise FileNotFoundError(f"TC template not found at {self.template_path}")
            else:
                logger.info("No TC template path provided. A default TC will be created.")
                doc = self._create_default_tc_template()

            self._replace_placeholders(doc, student_data, tc_data)

            # Generate a unique and predictable filename
            clean_adm_no = self._get_safe_filename_part(student_data.get('admission_no', ''))
            tc_number_safe = self._get_safe_filename_part(tc_data.get('tc_number', ''))
            base_filename = f"TC_{clean_adm_no}_{tc_number_safe}"
            
            docx_path = os.path.join(self.output_path_base, f"{base_filename}.docx")
            pdf_path = os.path.join(self.output_path_base, f"{base_filename}.pdf")

            doc.save(docx_path)
            logger.info(f"TC Word document generated successfully: {docx_path}")

            # Step 2: Convert Word Document to PDF
            try:
                pythoncom.CoInitialize() # Initialize COM
                convert(docx_path, pdf_path)
                logger.info(f"TC PDF generated successfully: {pdf_path}")
            finally:
                pythoncom.CoUninitialize() # Uninitialize COM


            return docx_path, pdf_path

        except Exception as e:
            logger.error(f"Error during TC file generation (from DOCX: {docx_path}): {e}", exc_info=True)
            # Clean up docx if PDF fails and docx exists
            if docx_path and os.path.exists(docx_path):
                try:
                    os.remove(docx_path)
                    logger.info(f"Cleaned up intermediate DOCX file after PDF generation failure: {docx_path}")
                except OSError as rm_e:
                    logger.error(f"Failed to cleanup DOCX file {docx_path}: {rm_e}")
            raise RuntimeError(f"Failed to generate TC files: {e}")

    def _replace_placeholders(self, doc: Document, student_data: dict, tc_data: dict):
        """Replaces placeholder strings in the document with actual data."""
        replacements = self._prepare_replacement_data(student_data, tc_data)

        # Helper to perform replacement in paragraphs (in body, tables, headers, footers)
        def replace_text_in_paragraph(paragraph):
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    inline = paragraph.runs
                    # Replace text in each run to preserve formatting
                    for i in range(len(inline)):
                        if placeholder in inline[i].text:
                            text = inline[i].text.replace(placeholder, str(value))
                            inline[i].text = text

        # Iterate through all parts of the document
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph)
        for section in doc.sections:
            for header_paragraph in section.header.paragraphs:
                replace_text_in_paragraph(header_paragraph)
            for footer_paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(footer_paragraph)

    def _format_display_date(self, date_value: Union[str, datetime, date, None]) -> str:
        """Formats a date into the app's display format (e.g., DD/MM/YYYY)."""
        if not date_value:
            return 'N/A'
        try:
            if isinstance(date_value, str):
                date_obj = datetime.strptime(date_value, '%Y-%m-%d') # Dates from DB are YYYY-MM-DD
            else:
                date_obj = date_value
            return date_obj.strftime(current_app.config.get('DISPLAY_DATE_FORMAT', '%d/%m/%Y'))
        except (ValueError, TypeError):
            return str(date_value) # Fallback

    def _date_to_words(self, date_value: Union[str, datetime, date, None]) -> str:
        """Converts a date into words (e.g., "Fifteenth August Nineteen Ninety-Nine")."""
        if not date_value:
            return 'N/A'
        
        try:
            if isinstance(date_value, str):
                date_obj = datetime.strptime(date_value, '%Y-%m-%d').date() # Dates from DB are YYYY-MM-DD
            elif isinstance(date_value, datetime):
                date_obj = date_value.date()
            elif isinstance(date_value, date):
                date_obj = date_value
            else:
                return str(date_value) # Fallback

            day_words = num2words(date_obj.day, to='ordinal').title()
            month_words = date_obj.strftime("%B")
            
            # Year to words - simple approach for now
            year = date_obj.year
            if 1900 <= year <= 2099: # Common range for TCs
                year_part1 = num2words(year // 100, to='cardinal').title()
                year_part2_num = year % 100
                if year_part2_num == 0:
                    year_words = f"{year_part1} Hundred"
                else:
                    year_part2 = num2words(year_part2_num, to='cardinal').title()
                    year_words = f"{year_part1} {year_part2}"
            else: # Fallback for years outside typical range
                year_words = num2words(year, to='year').title()

            return f"{day_words} {month_words} {year_words}"
        except Exception as e:
            logger.error(f"Error converting date to words ('{date_value}'): {e}", exc_info=True)
            return str(date_value) # Fallback

    def _format_date_explicit(self, date_value: Union[str, datetime, date, None], fmt: str) -> str:
        """Formats a date using an explicit strftime format string, converting to uppercase."""
        # This is similar to _format_display_date but for specific, non-configurable formats.
        # It assumes input string dates are in 'YYYY-MM-DD' format (from DB/schema).
        if not date_value: return 'N/A'
        try:
            if isinstance(date_value, str):
                date_obj = datetime.strptime(date_value, '%Y-%m-%d').date() # Schema uses YYYY-MM-DD
            elif isinstance(date_value, datetime):
                date_obj = date_value.date()
            elif isinstance(date_value, date):
                date_obj = date_value
            else: return str(date_value)
            return date_obj.strftime(fmt).upper()
        except (ValueError, TypeError) as e:
            logger.warning(f"Could not format date '{date_value}' with explicit format '{fmt}': {e}")
            return str(date_value)

    def _prepare_replacement_data(self, student_data: dict, tc_data: dict) -> dict:
        """Prepares a dictionary of placeholder keys and their corresponding values."""
        # Use a helper to safely get data and provide a default
        def get_data(key, default=''):
            return student_data.get(key) or tc_data.get(key) or default

         # --- Pronoun and Salutation Logic ---
        gender = str(get_data('gender', 'other')).lower()
        salutation = "Kum." if gender == 'female' else "Sri."
        possessive_pronoun = "Her" if gender == 'female' else "His"

        # --- Placeholder Value Logic ---
        admission_no = get_data('admission_no', '0000')
        # [TC_SNO] is the last four digits of the Admission Number 
        tc_sno = admission_no[-4:]

        # Get dates directly from the relevant source for clarity
        student_dob = student_data.get('dob')
        tc_form_date_of_leaving = tc_data.get('date_of_leaving') # Date pupil left, from TC form
        tc_form_issue_date = tc_data.get('issue_date') # Date TC is issued, from TC form
        
        return {
            '[TC_SNO]': tc_sno,
            '[ADM_NO]': admission_no,
            # Combine student name and surname into a single placeholder
            # If surname is provided, it will be appended with a space.
            # Otherwise, only the student name will be used.
            '[FULL_STUDENT_NAME]': f"{str(get_data('surname'))} {str(get_data('student_name'))}" if get_data('surname') else str(get_data('student_name')),
            # The old placeholders are kept for backward compatibility if your template still uses them,
            # but it's recommended to update your template to use [FULL_STUDENT_NAME].
            '[STUDENT_NAME]': str(get_data('student_name')), # Kept for backward compatibility
            '[SURNAME]': str(get_data('surname')).title() if get_data('surname') else '', # Kept for backward compatibility
            '[FATHER_NAME]': str(get_data('father_name')),
            '[SUB_CAS]': str(get_data('sub_caste')).title() if get_data('sub_caste') else '',
            '[NATIONALITY]': str(get_data('nationality', 'Indian')).title(),
            '[RELIGION]': str(get_data('religion')).title(),
            '[CASTE]': str(get_data('caste')).title(),
            '[DOB]': self._format_display_date(student_dob),
            '[DOB_WORDS]': tc_data.get('dob_in_words', 'N/A'), # Use pre-calculated DOB in words
            '[COURSE_NAME_FULL]': str(get_data('course_full_name') or get_data('course_name')),
            '[COURSE_NAME]': str(get_data('course_name')),
            '[COURSE_CODE]': str(get_data('course_code')), # For the line with parentheses 
            '[DATE_OF_ADMISSION]': self._format_date_explicit(student_data.get('date_of_admission'), '%b-%Y'), # Format: JUL-2020
            '[PROMOTION_STATUS]': tc_data.get('promotion_status', 'N/A'), # From manual entry 
            '[DATE_OF_LEAVING]': self._format_date_explicit(tc_form_date_of_leaving, '%b %Y'), # Format: MAY 2022
            '[TC_ISSUE_DATE]': self._format_date_explicit(tc_form_issue_date, '%d-%b-%Y'), # Format: 20-MAR-2023
            '[ACADEMIC_PERIOD]': str(get_data('academic_year') or 'N/A'), # For Study Certificate, use 'N/A' if data is missing
            '[CONDUCT]': str(tc_data.get('conduct') or student_data.get('conduct') or 'Satisfactory').title(), # Prioritize TC form conduct
            '[SALUTATION]': salutation, # For "Sri / Kum." 
            '[POSSESSIVE_PRONOUN]': possessive_pronoun # For "His / Her" 
        }

    def _create_default_tc_template(self) -> Document:
        """Creates a basic, default TC template if no external template is provided."""
        doc = Document()
        # Set to legal size: 8.5 x 14 inches
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
        # Optionally set margins
        margin_size = Inches(0.75)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size

        doc.add_heading(current_app.config.get('COLLEGE_NAME'), 0)
        doc.add_heading('Transfer Certificate', level=1)
        doc.add_paragraph("This is to certify that [STUDENT_NAME] (Adm. No: [ADMISSION_NO]) was a student of this college.")
        # ... add more placeholders
        return doc


def generate_tc_number_for_student(student_id: int) -> str:
    """
    Generates a unique TC number for the given student.
    Format: TC/YYYY/XXXX (XXXX is a 4-digit sequential number for that year).
    """
    from models.db_pool import db_manager
    year_str = str(datetime.now().year)

    last_tc_for_year = db_manager.execute_query(
        "SELECT tc_number FROM transfer_certificates WHERE tc_number LIKE ? ORDER BY tc_number DESC LIMIT 1",
        (f"TC/{year_str}/%",),
        fetch_one=True
    )

    next_serial = 1
    if last_tc_for_year and last_tc_for_year['tc_number']:
        try:
            last_serial_str = last_tc_for_year['tc_number'].split('/')[-1]
            next_serial = int(last_serial_str) + 1
        except (IndexError, ValueError):
            logger.warning(f"Could not parse serial from TC number {last_tc_for_year['tc_number']}. Defaulting to 1.")

    return f"TC/{year_str}/{next_serial:04d}"

class ReportGenerator:
    """Generates various reports (e.g., Admission Register) as Word documents."""

    def __init__(self):
        self.output_path_base = current_app.config.get('TC_OUTPUT_PATH') # Reports can go to same generated files area
        if not self.output_path_base:
            logger.error("TC_OUTPUT_PATH (used for reports) is not configured.")
            raise ValueError("Report output path is not configured.")
        os.makedirs(self.output_path_base, exist_ok=True)

    def _add_standard_header_to_doc(self, doc: Document, report_title_text: str, filter_info_text: Optional[str] = None):
        """Adds a standardized college name and report title to the document's header."""
        header = doc.sections[0].header
        
        # Clear existing paragraphs in the header to avoid duplication if called multiple times
        # or if the document object was somehow pre-populated.
        if header.paragraphs:
            for p_idx in range(len(header.paragraphs) - 1, -1, -1): # Iterate backwards for safe removal
                p_element = header.paragraphs[p_idx]._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)
                # Also clear the paragraph from the list in python-docx's object model
                # This part can be tricky; direct manipulation of header.paragraphs list might be needed
                # For simplicity, we assume add_paragraph works on a clean slate or appends.
                # A more robust clear might involve header.clear() if available or rebuilding the header part.
                # However, python-docx doesn't have a simple header.clear().
                # The most reliable way is to ensure you are working with a fresh header part or
                # carefully manage existing paragraphs. For now, we'll add, assuming it's okay for new docs.

        college_name_config = current_app.config.get('COLLEGE_NAME', 'YOUR COLLEGE NAME')
        
        college_p = header.add_paragraph()
        college_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        college_run = college_p.add_run(college_name_config.upper())
        college_run.font.size = Pt(20) # Further increased font size for college name
        college_run.bold = True

        title_p = header.add_paragraph() 
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(report_title_text.upper())
        title_run.font.size = Pt(16) # Increased font size for report title
        title_run.bold = True

        # Add filter information to the header if provided
        if filter_info_text:
            filter_p = header.add_paragraph()
            filter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            filter_run = filter_p.add_run(filter_info_text)
            filter_run.italic = True
            filter_run.font.size = Pt(12) # Increased font size for filter information

    def _add_standard_footer_to_doc(self, doc: Document):
        """Adds a standard footer with page numbering to the document."""
        footer = doc.sections[0].footer
        # Clear existing paragraphs in the footer
        if footer.paragraphs:
            for p_idx in range(len(footer.paragraphs) - 1, -1, -1):
                p_element = footer.paragraphs[p_idx]._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)

        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Align page number to the right
        
        # Add page number field (requires XML manipulation, but this is a standard pattern)
        # This inserts a Word field for PAGE number
        run = p.add_run()
        fldChar_begin = docx.oxml.OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = docx.oxml.OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve') # Important for Word to preserve space
        instrText.text = 'PAGE'
        fldChar_separate = docx.oxml.OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        # Add a default run for the page number, Word will update this.
        # run_default_text = docx.oxml.OxmlElement('w:r') 
        # t_default_text = docx.oxml.OxmlElement('w:t')
        # t_default_text.text = '1' # Default text
        # run_default_text.append(t_default_text)
        fldChar_end = docx.oxml.OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_separate)
        # run._r.append(run_default_text) # Word will insert the actual page number
        run._r.append(fldChar_end)

        # Apply font size to footer text
        for run_in_footer in p.runs: # Renamed 'run' to 'run_in_footer' to avoid conflict
            run_in_footer.font.size = Pt(9) # Smaller font for footer

    def _format_report_date(self, date_value):
        """Helper to format dates for reports."""
        if not date_value: return ''
        try:
            if isinstance(date_value, str):
                date_obj = datetime.strptime(date_value, current_app.config.get('DATE_FORMAT', '%Y-%m-%d')).date()
            elif isinstance(date_value, datetime):
                date_obj = date_value.date()
            elif isinstance(date_value, date):
                date_obj = date_value
            else: return str(date_value)
            return date_obj.strftime(current_app.config.get('DISPLAY_DATE_FORMAT', '%d/%m/%Y'))
        except ValueError:
            return str(date_value)

    def _add_text_to_cell(self, cell, text, bold=False, size=Pt(12), space_after=Pt(2), prefix="", font_name='Times New Roman'):
        """Helper to add formatted text to a cell, managing paragraphs."""
        # If text is explicitly an empty string, keep it empty. 
        # Otherwise, if it's None or whitespace-only (and not explicitly empty string), default to "N/A".
        if text == "": # Explicitly empty string
            text_to_add = ""
        elif text is None or str(text).strip() == "": # None or whitespace only
            text_to_add = "N/A" # Default for other cases if no prefix
        else:
            text_to_add = str(text)
        
        full_text = f"{prefix}{text_to_add}" if prefix else text_to_add

        # Always add a new paragraph for simplicity and consistent spacing control.
        if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip() and not cell.paragraphs[0].runs:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()

        run = p.add_run(full_text)
        run.font.name = font_name
        run.font.size = size
        run.bold = bold
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = space_after
        p.paragraph_format.line_spacing = 1.0

    def _set_cell_background_color(self, cell, color_hex: str):
        """Sets the background color of a table cell.

        Args:
            cell: The docx.table._Cell object.
            color_hex (str): Hexadecimal color code (e.g., "FFFF00" for yellow).
        """
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.first_child_found_in("w:shd")
        if shd is None:
            shd = docx.oxml.OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:fill"), color_hex)

    def generate_admission_register_pdf(self, course_id: Optional[int] = None, academic_year_id: Optional[int] = None) -> str:
        """
        Generates an Admission Register report as a PDF document.
        
        Args:
            course_id (Optional[int]): Filter by course ID.
            academic_year_id (Optional[int]): Filter by academic year ID.
        
        Returns:
            str: Path to the generated .pdf report.
        """
        from models.db_pool import get_students_for_admission_register

        # Fetch student data
        students = get_students_for_admission_register(course_id=course_id, academic_year_id=academic_year_id)

        # --- Document and Page Setup ---
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Custom page size: 14.67 inches width, 11.33 inches height
        section.page_width = Inches(18.67) # Set new page width
        section.page_height = Inches(11.33) # Set new page height (swapped for landscape)
        
        margin_size = Inches(0.5)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size
        # --- End Page Setup ---
        
        # Add the repeating header to every page
        # --- Add filter information to header ---
        filter_texts = []
        if course_id:
            # Fetch course name for display if possible
            from models.db_pool import db_manager
            course_info = db_manager.execute_query(
                "SELECT course_name, course_code, type FROM courses WHERE id = ?", 
                (course_id,), 
                fetch_one=True
            )
            if course_info:
                filter_texts.append(f"Course: {course_info['course_name']} ({course_info['course_code']}) - Type: {course_info['type']}")
            else:
                filter_texts.append(f"Course ID: {course_id} (Details not found)")

        if academic_year_id:
            from models.db_pool import db_manager
            ay_info = db_manager.execute_query("SELECT academic_year FROM academic_years WHERE id = ?", (academic_year_id,), fetch_one=True)
            filter_texts.append(f"Academic Year: {ay_info['academic_year']}" if ay_info else f"Academic Year ID: {academic_year_id}")

        if filter_texts:
            filter_info_str = "Filters Applied: " + "; ".join(filter_texts)
        else:
            filter_info_str = None

        # Add the repeating header to every page, including filter info
        self._add_standard_header_to_doc(doc, 'ADMISSION REGISTER', filter_info_str)
        doc.add_paragraph() # Add a spacer in the body before the table
        # --- End filter information ---

        # Add the repeating footer with page numbers
        self._add_standard_footer_to_doc(doc)
        if not students:
            doc.add_paragraph("No students found matching the criteria.")
        else:
            # Main table with 6 columns as per new layout
            table = doc.add_table(rows=1, cols=9) # Increased to 7 columns
            table.style = 'Table Grid' # Apply a table style
            table.autofit = False 
            table.allow_autofit = False

            # Define column widths (approximate, adjust as needed for A3 landscape)
            # Available width: 19.56 - 0.5 - 0.5 = 18.56 inches
            # Usable width: 18.67 - 1.0 = 17.67 inches. Reallocating space for Remarks and Adm. No. font size.
            col_widths = [ 
                Inches(0.4),  # S.No (Reduced)
                Inches(0.8),  # Adm. No
                Inches(3.5),  # Student Details (Increased)
                Inches(1.2),  # Social Category
                Inches(1.5),  # Education/Dates
                Inches(3.0),  # TC/Adm Date
                Inches(4.2)   # Remarks (Reduced to compensate)
            ]
            for i, width in enumerate(col_widths):
                if i < len(table.columns):
                    table.columns[i].width = width # Apply the calculated width

            # Header row
            # Using more concise headers for better fit
            hdr_cells = table.rows[0].cells
            # Slightly adjusted header strings for potentially better wrapping
            headers = [
                'S.No', 'Adm. No', 'Student Details', 'Social Category', 'Education/Dates', 'TC/Adm Date', 'Remarks'
            ]
            for i, header_text in enumerate(headers):
                cell = hdr_cells[i]
                # The revised _add_text_to_cell will handle using the cell's existing paragraph correctly.
                self._add_text_to_cell(cell, header_text, bold=True, size=Pt(15), space_after=Pt(4)) 
                # Set background color for header cells
                self._set_cell_background_color(cell, "D3D3D3") 

            # Set the first row as a repeating header row
            table.rows[0].header = True

            # --- Configuration for page breaks ---
            # Aim for up to 4 students per page on the large paper.
            max_rows_per_page = 4 

            # Data rows
            for idx, student_row_data in enumerate(students, 1):
                student = dict(student_row_data) # Convert sqlite3.Row to dict
                row_cells = table.add_row().cells
                
                # Check if a page break is needed before this row (idx is 1-based for students)
                if idx > 1 and (idx - 1) % max_rows_per_page == 0:
                    # Set page_break_before on the first paragraph of the first cell of this new row
                    row_cells[0].paragraphs[0].paragraph_format.page_break_before = True

                # Set minimum row height for data rows
                table.rows[idx].height = Inches(1.8) # Increased minimum row height for more vertical space
                table.rows[idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST # Allow row to expand if needed

                # Helper to clear default paragraph from cell before adding custom ones
                # This helper is no longer needed as _add_text_to_cell is more robust.

                # Cell 0: S.No
                self._add_text_to_cell(row_cells[0], str(idx))

                # Cell 1: Admission Number
                self._add_text_to_cell(row_cells[1], student.get('admission_no'), bold=True, size=Pt(14))

                # Cell 2: Student Details (stacked)
                # Per your description: Student Name, Father Name, Address, Phone No.
                student_name_text = student.get('student_name', 'N/A') # Default if None

                # Manually add "Name: " (normal) and student_name (bold) to the first paragraph of the cell
                name_paragraph = row_cells[2].paragraphs[0]
                # Clear any existing content from the default paragraph if it's truly empty
                if not (len(row_cells[2].paragraphs) == 1 and not name_paragraph.text.strip() and not name_paragraph.runs):
                    name_paragraph = row_cells[2].add_paragraph() # Should not happen for a fresh cell
                else: # Clear existing runs from the first default paragraph
                    for run_idx in range(len(name_paragraph.runs)-1, -1, -1):
                        p_run = name_paragraph.runs[run_idx]
                        name_paragraph._p.remove(p_run._r)

                name_paragraph.paragraph_format.space_before = Pt(0)
                name_paragraph.paragraph_format.space_after = Pt(2) # Default space for a line
                name_paragraph.paragraph_format.line_spacing = 1.0

                run_name_prefix = name_paragraph.add_run("Name: ")
                run_name_prefix.font.name = 'Times New Roman' # Font name
                run_name_prefix.font.size = Pt(12) # Font size
                run_name_prefix.bold = False
                run_name_value = name_paragraph.add_run(student_name_text)
                run_name_value.font.name = 'Times New Roman' # Font name
                run_name_value.font.size = Pt(12) # Font size
                run_name_value.bold = True

                # Subsequent details will use _add_text_to_cell, adding new paragraphs
                self._add_text_to_cell(row_cells[2], student.get('father_name'), prefix="Father: ", size=Pt(12)) # Father Name
                address_parts = [student.get(key) for key in ['address1', 'address2', 'address3', 'town'] if student.get(key)]
                address_str = ', '.join(filter(None, address_parts))
                # Add space after the address
                self._add_text_to_cell(row_cells[2], address_str, prefix="Addr: ", size=Pt(12)) # Address
                # Add more space after the phone number (last item in this cell's main content)
                self._add_text_to_cell(row_cells[2], student.get('phone_no'), prefix="Phone: ", size=Pt(12)) # Phone Number
                self._add_text_to_cell(row_cells[2], student.get('aadhar_no'), prefix="Aadhar: ", size=Pt(12)) # Aadhar Number

                # Cell 3: Social Category (stacked)
                # Per your description: Caste, Sub-Caste, Religion
                self._add_text_to_cell(row_cells[3], student.get('caste'), prefix="Caste: ")
                self._add_text_to_cell(row_cells[3], student.get('sub_caste'), prefix="Sub-Caste: ")
                self._add_text_to_cell(row_cells[3], student.get('religion'), prefix="Religion: ")

                # Cell 4: Education & Dates (stacked)
                # Per your description: DOB, previous clg studied, previous clg tc no/ date
                self._add_text_to_cell(row_cells[4], self._format_report_date(student.get('dob')), prefix="DOB: ")
                self._add_text_to_cell(row_cells[4], student.get('previous_college'), prefix="Prev. College: ")
                self._add_text_to_cell(row_cells[4], student.get('old_tc_no_date'), prefix="Prev. TC: ")

                # Cell 5: TC & Admission Date (stacked)
                self._add_text_to_cell(row_cells[5], self._format_report_date(student.get('date_of_admission')), prefix="Date of Adm: ")
                self._add_text_to_cell(row_cells[5], " ", prefix="Issued TC No: ")
                self._add_text_to_cell(row_cells[5], " ", prefix="TC Issue Date: ") # tc_issue_date from JOIN
                
                # Cell 6: Remarks (New separate column)
                # Ensure remarks_text is a string, defaulting to empty if None
                remarks_text = student.get('remarks') or '' 
                # Pass the stripped remarks. If empty, _add_text_to_cell will handle it as an empty string.
                self._add_text_to_cell(row_cells[6], remarks_text.strip())

                
                for cell_in_row in row_cells:
                    cell_in_row.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # Save document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"admission_register_{timestamp}"
        docx_filepath = os.path.join(self.output_path_base, f"{base_filename}.docx")
        pdf_filepath = os.path.join(self.output_path_base, f"{base_filename}.pdf")

        try:
            doc.save(docx_filepath)
            logger.info(f"Admission Register (DOCX) generated: {docx_filepath}")

            # Convert to PDF
            try:
                pythoncom.CoInitialize() # Initialize COM
                convert(docx_filepath, pdf_filepath)
                logger.info(f"Admission Register (PDF) generated: {pdf_filepath}")
            finally:
                pythoncom.CoUninitialize() # Uninitialize COM
            
            # Clean up the DOCX file after successful PDF conversion
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
                logger.info(f"Cleaned up intermediate DOCX file: {docx_filepath}")
            
            return pdf_filepath
        except Exception as e:
            logger.error(f"Error during Admission Register PDF generation: {e}", exc_info=True)
            raise RuntimeError(f"Failed to generate Admission Register PDF: {e}")

    def generate_fee_collection_report_pdf(self, start_date: Optional[str] = None, end_date: Optional[str] = None,
                                           course_id: Optional[int] = None, academic_year_id: Optional[int] = None) -> str:
        """Generates a Fee Collection report as a PDF document."""
        from models.db_pool import get_fee_payments_for_report, db_manager

        # Call get_fee_payments_for_report with only the arguments it expects
        # The start_date and end_date parameters in this method are no longer
        # used by get_fee_payments_for_report as per its definition in db_pool.py
        payments = get_fee_payments_for_report(course_id=course_id, academic_year_id=academic_year_id)


        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(18.67)
        section.page_height = Inches(11.33)
        margin_size = Inches(0.5)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size

        filter_texts = []
        if start_date: filter_texts.append(f"From: {self._format_report_date(start_date)}")
        if end_date: filter_texts.append(f"To: {self._format_report_date(end_date)}")
        if course_id:
            course_info = db_manager.execute_query(
                "SELECT course_name, course_code, type FROM courses WHERE id = ?", 
                (course_id,), 
                fetch_one=True
            )
            if course_info: 
                course_name_val = course_info['course_name'] if 'course_name' in course_info.keys() and course_info['course_name'] else ''
                course_code_val = course_info['course_code'] if 'course_code' in course_info.keys() and course_info['course_code'] else ''
                course_type_val = course_info['type'] if 'type' in course_info.keys() and course_info['type'] else ''
                filter_texts.append(f"Course: {course_name_val} ({course_code_val}) - Type: {course_type_val}")

        if academic_year_id:
            ay_info = db_manager.execute_query("SELECT academic_year FROM academic_years WHERE id = ?", (academic_year_id,), fetch_one=True)
            if ay_info and 'academic_year' in ay_info.keys() and ay_info['academic_year']:
                filter_texts.append(f"Academic Year: {ay_info['academic_year']}")
            elif ay_info: # ay_info exists but academic_year might be missing or None
                filter_texts.append(f"Academic Year ID: {academic_year_id} (Details not found or empty)")
        
        filter_info_str = f"Filters: {'; '.join(filter_texts)}" if filter_texts else None

        # Add the repeating header to every page
        self._add_standard_header_to_doc(doc, 'FEE COLLECTION REPORT', filter_info_str)
        doc.add_paragraph()

        if not payments:
            doc.add_paragraph("No fee payments found matching the criteria.")
        else:
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            headers = ['S.No', 'Payment Date', 'Adm. No', 'Student Name', 'Course', 'Academic Year', 'Amount Paid', 'Method', 'Transaction ID']
            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                self._add_text_to_cell(hdr_cells[i], header_text, bold=True, size=Pt(12))
            
            # Set the first row as a repeating header row
            table.rows[0].header = True

            total_collected = 0
            # Define column widths for Fee Collection Report (8 columns, 17.67 inches usable width). Target: ~16.4 inches.
            fee_col_widths = [
                Inches(0.5),   # S.No
                Inches(1.0),   # Payment Date
                Inches(1.0),   # Adm. No
                Inches(3.0),   # Student Name
                Inches(2.0),   # Course
                Inches(1.5),   # Academic Year
                Inches(1.5),   # Amount Paid
                Inches(1.5),   # Method
                Inches(2.0)    # Transaction ID
            ]
            for idx, p_data in enumerate(payments, 1):
                p_data = dict(p_data) # Ensure it's a dict
                row_cells = table.add_row().cells
                self._add_text_to_cell(row_cells[0], str(idx))
                self._add_text_to_cell(row_cells[1], self._format_report_date(p_data.get('payment_date')))
                self._add_text_to_cell(row_cells[2], p_data.get('admission_no'))
                self._add_text_to_cell(row_cells[3], p_data.get('student_name'))
                self._add_text_to_cell(row_cells[4], p_data.get('course_name'))
                self._add_text_to_cell(row_cells[5], p_data.get('academic_year'))
                amount_paid = p_data.get('amount_paid', 0)
                self._add_text_to_cell(row_cells[6], f"{amount_paid:.2f}")
                self._add_text_to_cell(row_cells[7], p_data.get('payment_method'))
                self._add_text_to_cell(row_cells[8], p_data.get('transaction_id'))
                total_collected += amount_paid

            # Apply column widths
            for i, width in enumerate(fee_col_widths):
                table.columns[i].width = width
            
            # Add total row
            total_row_cells = table.add_row().cells
            # Merge cells for "Total Collected" label
            merged_cell = total_row_cells[0].merge(total_row_cells[5]) # Merge first 6 cells
            self._add_text_to_cell(merged_cell, "Total Collected:", bold=True, size=Pt(12))
            # Place the total collected value in the 'Amount Paid' column (index 6)
            self._add_text_to_cell(total_row_cells[6], f"{total_collected:.2f}", bold=True, size=Pt(12))

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"fee_collection_report_{timestamp}"
        docx_filepath = os.path.join(self.output_path_base, f"{base_filename}.docx")
        pdf_filepath = os.path.join(self.output_path_base, f"{base_filename}.pdf")

        try:
            doc.save(docx_filepath)
            try:
                pythoncom.CoInitialize()
                convert(docx_filepath, pdf_filepath)
            finally:
                pythoncom.CoUninitialize()
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
            return pdf_filepath
        except Exception as e:
            logger.error(f"Error during PDF generation: {e}", exc_info=True)
            if os.path.exists(docx_filepath):
                try:
                    os.remove(docx_filepath)
                except Exception as cleanup_err:
                    logger.warning(f"Failed to cleanup docx: {cleanup_err}")
            raise RuntimeError(f"Failed to generate PDF: {e}")

    # Consider applying similar try/finally for CoInitialize/CoUninitialize
    # in generate_tc_issued_report_pdf and generate_student_fee_history_pdf

    def generate_tc_issued_report_pdf(self, course_id: Optional[int] = None, 
                                      academic_year_id: Optional[int] = None) -> str:
        """Generates a TC Issued report as a PDF document."""
        from models.db_pool import get_tc_issued_for_report, db_manager # Ensure db_manager is imported if not already

        tcs_issued = get_tc_issued_for_report(course_id=course_id, academic_year_id=academic_year_id)

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(18.67)
        section.page_height = Inches(11.33)
        margin_size = Inches(0.5)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size

        filter_texts = []
        if course_id:
            course_info = db_manager.execute_query("SELECT course_name FROM courses WHERE id = ?", (course_id,), fetch_one=True)
            if course_info: filter_texts.append(f"Course: {course_info['course_name']}")
        if academic_year_id:
            ay_info = db_manager.execute_query("SELECT academic_year FROM academic_years WHERE id = ?", (academic_year_id,), fetch_one=True)
            if ay_info: filter_texts.append(f"Academic Year: {ay_info['academic_year']}")

        filter_info_str = f"Filters: {'; '.join(filter_texts)}" if filter_texts else None
        self._add_standard_header_to_doc(doc, 'TC ISSUED REPORT', filter_info_str)
        doc.add_paragraph()

        if not tcs_issued:
            doc.add_paragraph("No TCs found matching the criteria.")
        else:
            table = doc.add_table(rows=1, cols=8) # Changed cols from 7 to 8 to match headers
            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            headers = ['S.No', 'TC Number', 'Issue Date', 'Student Name', 'Adm. No', 'Course', 'Date of Admission', 'Date of Leaving']
            hdr_cells = table.rows[0].cells
            # Define column widths for TC Issued Report (7 columns, 17.67 inches usable width). Target: ~17.1 inches.
            # Target: ~16.4 inches.
            tc_col_widths = [
                Inches(0.5),   # S.No
                Inches(2.0),   # TC Number
                Inches(2.0),   # Issue Date
                Inches(3.0),   # Student Name
                Inches(1.5),   # Adm. No
                Inches(2.5),   # Course
                Inches(1.5),   # Date of Admission (New)
                Inches(2.5)    # Date of Leaving
            ]
            for i, header_text in enumerate(headers): 
                self._add_text_to_cell(hdr_cells[i], header_text, bold=True, size=Pt(12))

            # Set the first row as a repeating header row
            table.rows[0].header = True

            for idx, tc_data in enumerate(tcs_issued, 1):
                tc_data = dict(tc_data) # Ensure it's a dict
                row_cells = table.add_row().cells
                self._add_text_to_cell(row_cells[0], str(idx))
                self._add_text_to_cell(row_cells[1], tc_data.get('tc_number'))
                self._add_text_to_cell(row_cells[2], self._format_report_date(tc_data.get('issue_date')))
                self._add_text_to_cell(row_cells[3], tc_data.get('student_name'))
                self._add_text_to_cell(row_cells[4], tc_data.get('admission_no'))
                self._add_text_to_cell(row_cells[5], tc_data.get('course_name'))
                self._add_text_to_cell(row_cells[6], self._format_report_date(tc_data.get('date_of_admission'))) # New column
                self._add_text_to_cell(row_cells[7], self._format_report_date(tc_data.get('date_of_leaving'))) # Moved to correct column
            
            # Apply column widths
            for i, width in enumerate(tc_col_widths):
                table.columns[i].width = width

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"tc_issued_report_{timestamp}"
        docx_filepath = os.path.join(self.output_path_base, f"{base_filename}.docx")
        pdf_filepath = os.path.join(self.output_path_base, f"{base_filename}.pdf")
        try:
            doc.save(docx_filepath)
            try:
                pythoncom.CoInitialize()
                convert(docx_filepath, pdf_filepath)
            finally:
                pythoncom.CoUninitialize()
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
            return pdf_filepath
        except Exception as e:
            logger.error(f"Error during TC Issued Report PDF generation: {e}", exc_info=True)
            raise RuntimeError(f"Failed to generate TC Issued Report PDF: {e}")

    def _add_spacer(self, doc: Document, size: Pt = Pt(6)):
        """Adds a small empty paragraph to act as a spacer."""
        p = doc.add_paragraph()
        run = p.add_run() # Add a run to apply font size
        run.font.size = size

    def generate_student_fee_history_pdf(self, student_data: dict, payments: list, summary: dict) -> str:
        """
        Generates a Fee History report for a single student as a PDF document.

        Args:
            student_data (dict): Dictionary containing student information.
            payments (list): List of fee payment dictionaries for the student.
            summary (dict): Dictionary containing fee summary (total_fee, total_paid, balance_due).

        Returns:
            str: Path to the generated .pdf report.
        """
        doc = Document()
        section = doc.sections[0]
        # Standard A4 Portrait
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69) # A4 height
        margin_size = Inches(0.75)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size

        # --- Header ---
        college_name_config = current_app.config.get('COLLEGE_NAME', 'YOUR COLLEGE NAME')
        p_college = doc.add_paragraph()
        p_college.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_college = p_college.add_run(college_name_config.upper())
        run_college.bold = True
        run_college.font.name = 'Arial'
        run_college.font.size = Pt(14)

        p_title = doc.add_heading('Student Fee History', level=1)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
        self._add_spacer(doc, Pt(6))

        # --- Student Information Section ---
        student_info_heading = doc.add_heading('Student Information', level=2)
        student_info_heading.runs[0].font.name = 'Arial'
        student_info_heading.runs[0].font.size = Pt(14)
        self._add_spacer(doc, Pt(3)) # Smaller spacer after heading

        info_table = doc.add_table(rows=4, cols=2)
        # Helper for this table
        def add_info_row(label, value):
            row_cells = info_table.add_row().cells
            self._add_text_to_cell(row_cells[0], label, bold=True, font_name='Calibri', size=Pt(12))
            self._add_text_to_cell(row_cells[1], value, font_name='Calibri', size=Pt(12))
        
        # Clear the initial row before adding new ones
        info_table._tbl.remove(info_table.rows[0]._tr)
        add_info_row('Student Name:', f"{student_data.get('student_name', '')} {student_data.get('surname', '')}")
        add_info_row('Admission No:', student_data.get('admission_no', ''))
        add_info_row('Academic Year:', student_data.get('academic_year', ''))
        self._add_spacer(doc, Pt(6))

        # --- Fee Payment Details Section ---
        payments_heading = doc.add_heading('Fee Payment Details', level=2)
        payments_heading.runs[0].font.name = 'Arial'
        payments_heading.runs[0].font.size = Pt(14)
        self._add_spacer(doc, Pt(3))

        if payments:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.autofit = False # Allow manual column widths
            table.allow_autofit = False

            headers = ['Payment Date', 'Amount Paid', 'Payment Method', 'Transaction ID', 'Remarks']
            hdr_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                self._add_text_to_cell(hdr_cells[i], header_text, bold=True, font_name='Calibri', size=Pt(12))

            for payment in payments:
                row_cells = table.add_row().cells
                self._add_text_to_cell(row_cells[0], self._format_report_date(payment.get('payment_date')), font_name='Calibri', size=Pt(12))
                self._add_text_to_cell(row_cells[1], f"{payment.get('amount_paid', 0):.2f}", font_name='Calibri', size=Pt(12))
                self._add_text_to_cell(row_cells[2], payment.get('payment_method', ''), font_name='Calibri', size=Pt(12))
                self._add_text_to_cell(row_cells[3], payment.get('transaction_id'), font_name='Calibri', size=Pt(12))
                self._add_text_to_cell(row_cells[4], payment.get('remarks'), font_name='Calibri', size=Pt(12))

            # Attempt to set column widths (approximate for A4 portrait)
            # Usable width: 8.27 - 0.75 - 0.75 = 6.77 inches
            widths = (Inches(1.2), Inches(1.0), Inches(1.3), Inches(1.5), Inches(1.7)) # Adjusted for 5 columns to fit ~6.7 inches
            for i, width in enumerate(widths):
                if i < len(table.columns): # Check if column index is valid
                    table.columns[i].width = width
        else:
            p_no_payments = doc.add_paragraph("No payments recorded for this student.")
            p_no_payments.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_no_payments.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.italic = True
        self._add_spacer(doc, Pt(6))

        # --- Fee Summary Section ---
        summary_heading = doc.add_heading('Fee Summary', level=2)
        summary_heading.runs[0].font.name = 'Arial'
        summary_heading.runs[0].font.size = Pt(14)
        self._add_spacer(doc, Pt(3))

        summary_table = doc.add_table(rows=3, cols=2)
        # Helper for summary table
        def add_summary_row(label, value, value_bold=False):
            row_cells = summary_table.add_row().cells
            self._add_text_to_cell(row_cells[0], label, bold=True, font_name='Calibri', size=Pt(12))
            self._add_text_to_cell(row_cells[1], value, bold=value_bold, font_name='Calibri', size=Pt(12))
        
        summary_table._tbl.remove(summary_table.rows[0]._tr)
        summary_table._tbl.remove(summary_table.rows[0]._tr)
        summary_table._tbl.remove(summary_table.rows[0]._tr)
        add_summary_row("Total Fee Applicable:", f"{summary.get('total_fee', 0):.2f}")
        add_summary_row("Total Amount Paid:", f"{summary.get('total_paid', 0):.2f}")
        add_summary_row("Balance Due:", f"{summary.get('balance_due', 0):.2f}", value_bold=True)
        self._add_spacer(doc, Pt(12)) # More space before footer

        # --- Footer ---
        p_footer = doc.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = p_footer.add_run(f"Report generated on: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
        run_footer.font.size = Pt(9)
        run_footer.italic = True

        base_filename = f"fee_history_{student_data.get('admission_no', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        docx_filepath = os.path.join(self.output_path_base, f"{base_filename}.docx")
        pdf_filepath = os.path.join(self.output_path_base, f"{base_filename}.pdf")

        doc.save(docx_filepath)
        try:
            pythoncom.CoInitialize()
            convert(docx_filepath, pdf_filepath)
        finally:
            pythoncom.CoUninitialize() # Ensure COM is uninitialized
        if os.path.exists(docx_filepath): os.remove(docx_filepath)
        return pdf_filepath

    def generate_fee_summary_report_pdf(self, course_id: Optional[int] = None, academic_year_id: Optional[int] = None) -> str:
        """
        Generates a Fee Summary report as a PDF document.
        
        Args:
            course_id (Optional[int]): Filter by course ID.
            academic_year_id (Optional[int]): Filter by academic year ID.
        
        Returns:
            str: Path to the generated .pdf report.
        """
        from models.db_pool import db_manager

        # Fetch summary data
        query = """
        SELECT
            s.id AS student_id,
            s.student_name,
            s.surname,
            s.admission_no,
            c.course_name,
            ay.academic_year,
            COALESCE(fs.total_fee, 0) AS total_fee,
            COALESCE(SUM(p.amount_paid), 0) AS total_paid
        FROM
            students s
        JOIN
            courses c ON s.course_id = c.id
        JOIN
            academic_years ay ON s.academic_year_id = ay.id
        LEFT JOIN
            fee_structure fs ON s.fee_structure_id = fs.id
        LEFT JOIN
            student_fee_payments p ON s.id = p.student_id
        WHERE 1=1
        """
        params = []

        if course_id is not None:
            query += " AND s.course_id = ?"
            params.append(course_id)
        if academic_year_id is not None:
            query += " AND s.academic_year_id = ?"
            params.append(academic_year_id)

        query += """ GROUP BY
            s.id, s.student_name, s.surname, s.admission_no, c.course_name, ay.academic_year, fs.total_fee
        ORDER BY
            c.course_name, ay.academic_year, s.student_name, s.surname;
        """
        summaries = db_manager.execute_query(query, params, fetch_all=True)
        
        # --- Document and Page Setup (similar to other reports) ---
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        # Custom page size: 14.67 inches width, 11.33 inches height
        section.page_width = Inches(18.67) # Set new page width
        section.page_height = Inches(11.33) # Set new page height (swapped for landscape)
        
        margin_size = Inches(0.5)
        section.top_margin = margin_size
        section.bottom_margin = margin_size
        section.left_margin = margin_size
        section.right_margin = margin_size
        # --- End Page Setup ---
        
        # Add the repeating header to every page
        # --- Add filter information to header ---
        filter_texts = []
        if course_id:
            # Fetch course name for display if possible
            from models.db_pool import db_manager
            course_info = db_manager.execute_query(
                "SELECT course_name, course_code, type FROM courses WHERE id = ?", 
                (course_id,), 
                fetch_one=True
            )
            if course_info:
                filter_texts.append(f"Course: {course_info['course_name']} ({course_info['course_code']}) - Type: {course_info['type']}")
            else:
                filter_texts.append(f"Course ID: {course_id} (Details not found)")

        if academic_year_id:
            from models.db_pool import db_manager
            ay_info = db_manager.execute_query("SELECT academic_year FROM academic_years WHERE id = ?", (academic_year_id,), fetch_one=True)
            filter_texts.append(f"Academic Year: {ay_info['academic_year']}" if ay_info else f"Academic Year ID: {academic_year_id}")

        if filter_texts:
            filter_info_str = "Filters Applied: " + "; ".join(filter_texts)
        else:
            filter_info_str = None

        # Add the repeating header to every page, including filter info
        self._add_standard_header_to_doc(doc, 'FEE SUMMARY REPORT', filter_info_str)
        doc.add_paragraph() # Add a spacer in the body before the table
        # --- End filter information ---

        # Add the repeating footer with page numbers
        self._add_standard_footer_to_doc(doc)
        if not summaries:
            doc.add_paragraph("No fee summary data found matching the criteria.")
        else:
            # Corrected table for Fee Summary Report
            table = doc.add_table(rows=1, cols=8) # S.No, Adm. No, Student Name, Course, Academic Year, Total Fee, Total Paid, Balance Due
            table.style = 'Table Grid' # Apply a table style
            table.autofit = False 
            table.allow_autofit = False

            # Define column widths for Fee Summary Report (8 columns, 17.67 inches usable width)
            col_widths = [ 
                Inches(0.5),   # S.No
                Inches(1.0),   # Adm. No
                Inches(3.0),   # Student Name
                Inches(2.0),   # Course
                Inches(1.5),   # Academic Year
                Inches(1.5),   # Total Fee
                Inches(1.5),   # Total Paid
                Inches(1.5)    # Balance Due
            ]
            for i, width in enumerate(col_widths):
                if i < len(table.columns):
                    table.columns[i].width = width # Apply the calculated width

            # Header row
            hdr_cells = table.rows[0].cells
            headers = [
                'S.No', 'Adm. No', 'Student Name', 'Course', 'Academic Year', 'Total Fee', 'Total Paid', 'Balance Due'
            ]
            for i, header_text in enumerate(headers):
                cell = hdr_cells[i]
                self._add_text_to_cell(cell, header_text, bold=True, size=Pt(12), space_after=Pt(4)) 
                # Set background color for header cells
                self._set_cell_background_color(cell, "D3D3D3") 

            # Set the first row as a repeating header row
            table.rows[0].header = True

            # Data rows
            total_fees_due = 0.0
            total_collected = 0.0
            total_fee_applicable = 0.0
            for idx, summary_row_data in enumerate(summaries, 1):
                summary = dict(summary_row_data) # Convert sqlite3.Row to dict
                row_cells = table.add_row().cells
                
                # Cell 0: S.No
                self._add_text_to_cell(row_cells[0], str(idx))
                # Cell 1: Admission Number
                self._add_text_to_cell(row_cells[1], summary.get('admission_no'))
                # Cell 2: Student Name
                self._add_text_to_cell(row_cells[2], f"{summary.get('student_name', '')} {summary.get('surname', '')}".strip())
                # Cell 3: Course
                self._add_text_to_cell(row_cells[3], summary.get('course_name'))
                # Cell 4: Academic Year
                self._add_text_to_cell(row_cells[4], summary.get('academic_year'))
                # Cell 5: Total Fee
                total_fee = summary.get('total_fee', 0)
                self._add_text_to_cell(row_cells[5], f"{total_fee:.2f}")
                total_fee_applicable += total_fee
                # Cell 6: Total Paid
                total_paid = summary.get('total_paid', 0)
                self._add_text_to_cell(row_cells[6], f"{total_paid:.2f}")
                total_collected += total_paid
                # Cell 7: Balance Due
                balance_due = total_fee - total_paid
                self._add_text_to_cell(row_cells[7], f"{balance_due:.2f}")
                total_fees_due += balance_due

                for cell_in_row in row_cells:
                    cell_in_row.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
            # Add total rows at the end of the table
            # Total Fee Applicable
            total_fee_row = table.add_row().cells
            merged_cell_fee = total_fee_row[0].merge(total_fee_row[4])
            self._add_text_to_cell(merged_cell_fee, "Total Fee Applicable:", bold=True, size=Pt(12))
            self._add_text_to_cell(total_fee_row[5], f"{total_fee_applicable:.2f}", bold=True, size=Pt(12))
            # Total Collected
            total_collected_row = table.add_row().cells
            merged_cell_collected = total_collected_row[0].merge(total_collected_row[5])
            self._add_text_to_cell(merged_cell_collected, "Total Collected:", bold=True, size=Pt(12))
            self._add_text_to_cell(total_collected_row[6], f"{total_collected:.2f}", bold=True, size=Pt(12))
            # Total Fees Due
            total_due_row = table.add_row().cells
            merged_cell_due = total_due_row[0].merge(total_due_row[6])
            self._add_text_to_cell(merged_cell_due, "Total Fees Due:", bold=True, size=Pt(12))
            self._add_text_to_cell(total_due_row[7], f"{total_fees_due:.2f}", bold=True, size=Pt(12))

        # Save document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"fee_summary_report_{timestamp}"
        docx_filepath = os.path.join(self.output_path_base, f"{base_filename}.docx")
        pdf_filepath = os.path.join(self.output_path_base, f"{base_filename}.pdf")

        try:
            doc.save(docx_filepath)
            try:
                pythoncom.CoInitialize()
                convert(docx_filepath, pdf_filepath)
            finally:
                pythoncom.CoUninitialize()
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
            return pdf_filepath
        except Exception as e:
            logger.error(f"Error during PDF generation: {e}", exc_info=True)
            if os.path.exists(docx_filepath):
                try:
                    os.remove(docx_filepath)
                except Exception as cleanup_err:
                    logger.warning(f"Failed to cleanup docx: {cleanup_err}")
            raise RuntimeError(f"Failed to generate PDF: {e}")
